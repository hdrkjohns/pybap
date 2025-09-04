import datetime as dt
#from arcgis.gis import GIS
import stat
import sys
import os
import io
import zipfile
import shutil
import time
import pkg_resources
from osgeo import gdal, ogr
import geopandas as gpd
import pandas as pd
import numpy as np
import openpyxl as xl
import pandas as pd
import numpy as np
from shapely.geometry import Point
from docx import Document
from pyproj import Transformer

FS_ITEM_ID = 'd92220452d9346d7964b41ed0103806c'
FS_URL = 'https://services.arcgis.com/04HiymDgLlsbhaV4/arcgis/rest/services/BAP3/FeatureServer'

if(sys.platform == 'win32'):
    OUT_DIR = r'C:\Users\kjohns\OneDrive - HDR, Inc\Documents\BAP\_EXPORT'
    EXCEL_BAP_TEMPLATE = r"C:\Users\kjohns\HDR, Inc\Gambrell, Travis - Templates\(Proj Name) - Bldg Assessment Form.xlsx"
    WORD_BAP_TEMPLATE = r"C:\Users\kjohns\HDR, Inc\Gambrell, Travis - Templates\(Proj Name) - Cover Page and TOC.docx"
    EXCEL_BAP_COMPONENT_TEMPLATE = r"C:\Users\kjohns\HDR, Inc\Gambrell, Travis - Templates\(Proj Name) - Major Component List - Simple.xlsx"
elif(sys.platform == 'linux'):
    OUT_DIR = r'/mnt/rtc/bap/exports'
    EXCEL_BAP_TEMPLATE = r"/mnt/rtc/bap/templates/(Proj Name) - Bldg Assessment Form.xlsx"
    WORD_BAP_TEMPLATE = r"/mnt/rtc/bap/templates/(Proj Name) - Cover Page and TOC.docx"
    EXCEL_BAP_COMPONENT_TEMPLATE = r"/mnt/rtc/bap/templates/(Proj Name) - Major Component List - Simple.xlsx"

OUT_ZIP_BASE = 'BAP_EXPORT'

# EXCEL_COL_LOOKUP = r"C:\Users\kjohns\OneDrive - HDR, Inc\Documents\BAP\BAP_Excel_Vals.xlsx"

LU_CATEGORY_CODES = {'N/A': 'N/A', 'POOR': 'Poor', 'FAIR': 'Fair', 'CRIT': 'Safety / Critical', 'SFTY': 'Safety / Critical', 'ACT': 'Acceptable', 'GWO': 'Acceptable'}

lookup_excel_path = pkg_resources.resource_filename('pybap', 'assets/BAP_Excel_Vals.xlsx')

df_lookup = pd.read_excel(lookup_excel_path, sheet_name='BAP_Excel_Vals')

def download_online_gdb(fs_item_id:str, out_zip_base:str, out_dir:str) -> str:
    '''
    '''

    ## what the file geodatabase will be called in ArcGIS Online
    export_title = f"{out_zip_base}_{dt.datetime.now().strftime('%Y%m%d_%H%M')}"

    ## the export format
    export_format = "File Geodatabase"

    ## where to save the download
    save_path = out_dir

    ## access agol
    agol = GIS("home")

    ## access the feature service
    item = agol.content.get(fs_item_id)

    ## when we call the export() method from an Item object we are returned
    ## the Item object from the export
    export = item.export(
        title = export_title,
        export_format = export_format
    )

    ## use the download() method from the export Item object to download the
    ## file geodatabase in zipped folder
    export.download(
        save_path = save_path
    )

    export.delete(force=True, permanent=True)

    out_zip_file = os.path.join(out_dir, f'{export_title}.zip')

    return out_zip_file

def unzip_and_rename(out_zip_file:str, out_dir:str):
    '''
    '''
    def del_rw(action, name, exc):
        print(exc)
        os.chmod(name, stat.S_IWRITE)
        os.remove(name)

    # unzip it now, and then rename the .gdb
    z = zipfile.ZipFile(out_zip_file, mode='r')
    dirs = list(set([os.path.dirname(x) for x in z.namelist()]))

    z.extractall(out_dir)
    z.close()

    if(len(dirs) == 1):
        gdb_path = os.path.join(out_dir, 'BAP.gdb')

        # delete if exists
        if(os.path.exists(gdb_path)):
            shutil.rmtree(gdb_path, onerror=del_rw)

        # sleep
        time.sleep(1)

        # rename
        os.rename(os.path.join(out_dir, dirs[0]), gdb_path)

        return gdb_path

    return None

def bap_gdb_to_dataframe(gdb_file) -> gpd.GeoDataFrame:
    '''
    '''
    ds = gdal.OpenEx(gdb_file, gdal.OF_VECTOR, open_options=['LIST_ALL_TABLES=YES'])
    layer_names = {ds.GetLayer(i).GetName():i for i in range(ds.GetLayerCount())}

    df_main = gpd.read_file(gdb_file, layer='Asset_Points')
    df_main.columns += '_main'

    df_ai = gpd.read_file(gdb_file, layer='Architecture_Interior')
    df_ai.columns += '_ai'

    df_ae = gpd.read_file(gdb_file, layer='Architecture_Exterior')
    df_ae.columns += '_ae'

    df_el = gpd.read_file(gdb_file, layer='Electrical')
    df_el.columns += '_el'

    df_hvac = gpd.read_file(gdb_file, layer='Mechanical_HVAC')
    df_hvac.columns += '_hvac'

    df_plumbing = gpd.read_file(gdb_file, layer='Mechanical_Plumbing')
    df_plumbing.columns += '_plumb'

    df_sec = gpd.read_file(gdb_file, layer='Security')
    df_sec.columns += '_sec'

    df_sc = gpd.read_file(gdb_file, layer='Site_Civil')
    df_sc.columns += '_sc'

    df_st = gpd.read_file(gdb_file, layer='Structural')
    df_st.columns += '_st'

    try:
        ds.Close()
    except AttributeError:
        pass
    finally:
        del ds

    # join all the things together to make the big table
    df_all = df_ai.merge(df_ae, how='inner', left_on='asset_point_globalid_ai', right_on='asset_point_globalid_ae') \
            .merge(df_el, how='inner', left_on='asset_point_globalid_ai', right_on='asset_point_globalid_el') \
            .merge(df_hvac, how='inner', left_on='asset_point_globalid_ai', right_on='asset_point_globalid_hvac') \
            .merge(df_plumbing, how='inner', left_on='asset_point_globalid_ai', right_on='asset_point_globalid_plumb') \
            .merge(df_sec, how='inner', left_on='asset_point_globalid_ai', right_on='asset_point_globalid_sec') \
            .merge(df_sc, how='inner', left_on='asset_point_globalid_ai', right_on='asset_point_globalid_sc') \
            .merge(df_st, how='inner', left_on='asset_point_globalid_ai', right_on='asset_point_globalid_st') \
            .merge(df_main, how='left', left_on='asset_point_globalid_ai', right_on='GlobalID_main')

    # make the globalid_main the index
    df_all.index = df_all['GlobalID_main']
    df_all.replace(np.nan, None, inplace=True)

    return df_all

def generate_bap_excel(gdf:gpd.GeoDataFrame, asset_guid:str, out_bap_report:str, out_bap_component:str):
    '''
    '''

    trans = Transformer.from_crs(
        "EPSG:3857",
        "EPSG:4326",
        always_xy=True,
    )

    #display(df_lookup)

    # assessment form
    wb = xl.open(EXCEL_BAP_TEMPLATE)

    # component list
    wb_component = xl.open(EXCEL_BAP_COMPONENT_TEMPLATE)
    sheet_component = wb_component['Combined Sheet']
    row_component = 2

    gdf_filter = gdf[gdf['GlobalID_main'] == asset_guid]
    gdf_filter_t = gdf_filter.T.replace(np.nan, None)
    #display(gdf_filter_t)

    # go through the row lookups
    for idx, lu_row in df_lookup.iterrows():
        sheet = wb[lu_row['sheet_name']]
        rownum = lu_row['rownum']
        colnum = lu_row['colnum']
        field_name = lu_row['field_name']
        include_in_components = lu_row['include_in_components'] == 1

        if(field_name and not (type(field_name) == float and np.isnan(field_name))):
            orig_field, table = str(field_name).rsplit('_', maxsplit=1)
            field_name_comment = f'{orig_field}_comments_{table}'

            if(field_name == 'geometry_main'):
                val = gdf_filter_t.loc[field_name][asset_guid]
                xx, yy = trans.transform([val.x], [val.y])

                cell_val = Point((yy[0], xx[0]))
            else:
                cell_val = gdf_filter_t.loc[field_name][asset_guid]
            #cell_comments = gdf_filter_t.loc[field_name][asset_guid]

            if(type(cell_val) == Point):
                cell_val = f'{cell_val.x:.2f}, {cell_val.y:.2f}'

            comments = '' if field_name_comment not in gdf_filter_t.index else gdf_filter_t.loc[field_name_comment][asset_guid]

            if(comments is not None and len(comments.strip()) > 0 and comments != ' None'):
                comments = ' - ' + comments
            else:
                comments = ''

            cell = sheet.cell(rownum, colnum)
            cell_text = '' if cell.value is None else str(cell.value.replace('_', '').strip())
            if(not cell_text.endswith(':')):
                cell_text += ':' if orig_field != 'comments' else ''

            cell.value = f'{cell_text} {cell_val}{comments}'

            # component sheet
            if(include_in_components):
                category = LU_CATEGORY_CODES.get(cell_val, 'N/A')

                if(category != 'N/A'):
                    cell_component = sheet_component.cell(row_component, 4)
                    cell_component.value = lu_row['sheet_name']

                    cell_component = sheet_component.cell(row_component, 5)
                    cell_component.value = cell_text[:-1]

                    cell_component = sheet_component.cell(row_component, 7)
                    cell_component.value = category

                    row_component += 1

    wb.save(out_bap_report)
    wb_component.save(out_bap_component)

    return (wb, wb_component)

def generate_bap_worddoc(gdf:gpd.GeoDataFrame, asset_guid:str, in_docx_template:str, out_docx:str) -> Document:
    '''
    '''

    gdf_filter = gdf[gdf['GlobalID_main'] == asset_guid]
    gdf_filter_t = gdf_filter.T.replace(np.nan, None)

    # just get poor, actionable safety or safet concerns
    gdf_filter_poor = gdf_filter_t[gdf_filter_t[asset_guid].isin(['POOR', 'CRIT', 'SFTY'])]

    # join to the lookup
    gdf_filter_poor = pd.merge(gdf_filter_poor, df_lookup, how='inner', left_index=True, right_on='field_name' )
    #display(gdf_filter_poor)

    # Create a Word document
    doc = Document()

    for sheet_name in gdf_filter_poor['sheet_name'].unique():
        doc.add_heading(sheet_name, level=1)

        for idx, elementrow in gdf_filter_poor[gdf_filter_poor['sheet_name'] == sheet_name].iterrows():
            doc.add_heading(f"{elementrow['val']} ({elementrow[asset_guid]})", level=3)

    # Add a table to the document
    if(False):
        table = doc.add_table(rows=1, cols=len(gdf_filter.columns))
        table.style = "Table Grid"

        # Add header row
        header_cells = table.rows[0].cells
        for i, column_name in enumerate(gdf_filter.columns):
            header_cells[i].text = column_name

        # Add DataFrame rows to the table
        for _, row in gdf_filter.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

    # Save the document
    doc.save(out_docx)

    return doc

def combine_bap_files(out_bap_report:str, out_bap_component:str, out_docx:str) -> io.BytesIO:
    '''
    '''
    buffer = io.BytesIO()
    z = zipfile.ZipFile(buffer, mode='w')

    z.write(out_bap_report, os.path.basename(out_bap_report))
    z.write(out_bap_component, os.path.basename(out_bap_component))
    z.write(out_docx, os.path.basename(out_docx))

    z.close()
    buffer.seek(0)

    return buffer